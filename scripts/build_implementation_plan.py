from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


WORKDIR = Path(__file__).resolve().parents[1]
OUT = WORKDIR / "docs" / "source" / "开店指导模拟预测软件_详细实施计划书.docx"
TABLE_HELPERS = Path(
    r"C:\Users\Administrator\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.601.10930\skills\documents\scripts"
)
sys.path.insert(0, str(TABLE_HELPERS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(24, 28, 35)
MUTED = RGBColor(88, 96, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF7E0"
PALE_GREEN = "EAF6EE"
PALE_RED = "FCEEEE"


def set_east_asia_font(run, east_asia="Microsoft YaHei", latin="Calibri"):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border_bottom(paragraph, color="D7DBE2", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_header_row(table):
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_run(paragraph, text, *, size=11, bold=False, italic=False, color=INK):
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def add_para(doc, text="", *, style=None, size=11, bold=False, color=INK, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    add_run(
        p,
        text,
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color=BLUE if level < 3 else DARK_BLUE,
    )
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_run(p, text)
    return p


def add_number(doc, text, level=0):
    p = doc.add_paragraph(style="List Number" if level == 0 else "List Number 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_run(p, text)
    return p


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Twips(120)
    p.paragraph_format.right_indent = Twips(120)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    set_paragraph_shading(p, fill)
    add_run(p, f"{title}：", bold=True, color=DARK_BLUE)
    add_run(p, body, color=INK)


def add_table(doc, headers, rows, widths, *, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_fill(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, header, size=font_size, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_run(p, str(value), size=font_size, color=INK)
    set_table_header_row(table)
    apply_table_geometry(table, widths, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Bullet 2", "List Number", "List Number 2"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(header, "开店指导模拟预测软件 | 详细实施计划书", size=9, color=MUTED)
    set_paragraph_border_bottom(header, color="D7DBE2", size="4")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "内部执行稿 | 按阶段复盘与迭代", size=9, color=MUTED)


def add_title_page(doc):
    add_para(doc, "实施计划书", size=12, bold=True, color=BLUE, after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "开店指导模拟预测软件", size=26, bold=True, color=RGBColor(0, 0, 0))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(20)
    add_run(p2, "从“位置+品类”可行性报告，到多智能体经营策略模拟的分步执行方案", size=13, color=MUTED)

    meta = [
        ("来源文档", "项目书.docx"),
        ("生成日期", "2026年6月10日"),
        ("建议路线", "先做8-10周MVP，再推进36周完整商业化路线"),
        ("执行原则", "每次只做一个可验收的小模块，避免一次性搭建复杂系统"),
    ]
    add_table(
        doc,
        ["项目", "内容"],
        meta,
        column_widths_from_weights([1.4, 5.1]),
        header_fill=LIGHT_GRAY,
        font_size=10,
    )

    add_callout(
        doc,
        "核心判断",
        "这个项目不要从“完整MiroFish集成+商业上线”直接开始。最稳妥的路线是先完成一个能输入地址和品类、能抓取周边POI、能计算基础评分、能生成可读报告的MVP；等闭环跑通后，再把MiroFish作为增强引擎接入。",
        fill=PALE_GREEN,
    )


def add_overview(doc):
    add_heading(doc, "一、项目目标拆解", 1)
    add_para(
        doc,
        "原项目书的目标可以拆成三层：第一层是开店前选址评估；第二层是已有店铺经营策略模拟；第三层是基于MiroFish思路的多智能体市场推演。实际执行时，必须先完成第一层，否则后两层没有稳定的数据输入和验证标准。",
    )
    add_table(
        doc,
        ["层级", "要实现的能力", "先后顺序", "判断标准"],
        [
            ("基础层", "用户输入地址/地图选点和开店品类，系统抓取周边POI并输出可行性报告", "必须最先做", "普通用户能在3分钟内得到一份有评分、有理由、有风险提示的报告"),
            ("增强层", "加入租金、人力、客单价、营业时间、外卖、优惠等变量，模拟经营策略变化", "MVP后做", "用户调整一个变量后，能看到收入、成本、风险的变化方向和解释"),
            ("差异化层", "接入或参考MiroFish，构建消费者、竞品、供应商等智能体进行市场推演", "数据稳定后做", "模拟结果能补充基础模型，而不是替代基础模型"),
            ("商业层", "账号体系、报告保存、付费套餐、合规备案、云部署和运营", "验证需求后做", "有真实用户愿意反复使用并为报告或咨询付费"),
        ],
        column_widths_from_weights([1.0, 2.4, 1.3, 1.8]),
    )

    add_heading(doc, "二、推荐技术路线", 1)
    add_para(doc, "为降低学习和调试成本，建议采用“前端页面 + Python后端 + 数据库 + LLM报告生成 + MiroFish独立验证”的分层架构。")
    add_table(
        doc,
        ["模块", "建议选择", "为什么这样选", "后续可升级"],
        [
            ("前端", "Next.js或普通React网页；小程序可放到后期", "网页调试快，适合先做原型和后台管理", "微信小程序、移动端适配"),
            ("后端", "FastAPI", "Python生态适合数据计算、地图API、机器学习和LLM调用", "拆成多个服务或任务队列"),
            ("数据库", "SQLite起步，PostgreSQL上线", "SQLite零配置，适合MVP；PostgreSQL适合正式部署", "PostGIS地理空间能力"),
            ("地图数据", "高德地图POI接口为主，OSM/Nominatim作备选", "高德适合国内地址与POI；备选能降低依赖", "商业数据商、品牌合作数据"),
            ("模型", "规则评分 + IRS + 哈夫模型起步", "可解释、容易调试，不依赖大量训练数据", "LightGBM/XGBoost、蒙特卡洛模拟"),
            ("报告", "Qwen/GLM/GPT等LLM生成自然语言报告", "把计算结果转成用户能理解的建议", "模板化报告、PDF导出、行业报告库"),
            ("MiroFish", "先本地跑通，再决定直接改造或独立重写", "避免早期被协议、环境和复杂度拖慢", "多智能体数字孪生引擎"),
        ],
        column_widths_from_weights([1.0, 1.5, 2.3, 1.7]),
    )

    add_heading(doc, "三、总路线图", 1)
    add_table(
        doc,
        ["阶段", "时间", "核心目标", "交付物", "是否依赖我后续协助"],
        [
            ("0. 准备", "1-3天", "确定MVP边界、装好工具、申请API Key", "项目文件夹、技术选型、账号清单", "是，可逐项带你完成"),
            ("1. 产品原型", "第1-2周", "画出用户流程和页面原型", "PRD、字段清单、低保真原型", "是，可直接生成原型"),
            ("2. 数据闭环", "第3-4周", "地址转经纬度，抓取周边POI并分类", "高德API采集脚本、POI分类结果", "是，可写代码并调试"),
            ("3. 基础模型", "第5-6周", "计算选址评分和风险因子", "评分引擎、可解释指标", "是，可逐个公式实现"),
            ("4. 报告生成", "第7-8周", "把评分结果生成可读报告", "报告页面、LLM提示词、导出文本", "是，可做提示词和页面"),
            ("5. MVP内测", "第9-10周", "找真实地址试用，修正模型参数", "测试记录、问题清单、MVP v0.1", "是，可分析测试结果"),
            ("6. MiroFish增强", "第11-18周", "验证并接入多智能体仿真", "本地部署记录、智能体输入格式、模拟报告", "是，可逐步研究接入"),
            ("7. 商业化准备", "第19-36周", "账号、支付、部署、合规、运营", "上线版本、隐私政策、运营看板", "是，可按模块推进"),
        ],
        column_widths_from_weights([1.0, 0.9, 2.0, 1.8, 0.8]),
        font_size=9,
    )


def add_phase_zero(doc):
    add_heading(doc, "四、阶段0：准备工作（1-3天）", 1)
    add_para(doc, "这一阶段不写复杂代码，目标是把后面会卡住的账号、环境、项目边界先整理好。")
    add_table(
        doc,
        ["步骤", "你要做什么", "我可以帮你做什么", "完成标准"],
        [
            ("0.1", "确定首个试点城市和首个品类，例如“广州奶茶店”或“成都美甲店”", "帮你把该品类的核心指标和评分权重列出来", "只选1个城市+1个品类，不同时做很多行业"),
            ("0.2", "申请高德开放平台账号并创建Web服务Key", "告诉你需要打开哪些API权限，并写测试脚本", "能用Key请求周边POI接口"),
            ("0.3", "准备LLM API Key，例如阿里云百炼Qwen、智谱GLM或OpenAI", "封装统一调用函数，避免后续换模型时大改代码", "能成功发起一次文本生成请求"),
            ("0.4", "安装开发工具：Git、VS Code/Cursor、Node.js、Python", "检查环境版本并初始化项目", "命令行能运行node、python、git"),
            ("0.5", "决定MiroFish策略：先研究架构，暂不把源码混入主项目", "帮你单独建一个实验目录跑MiroFish", "主项目和MiroFish实验目录分开"),
        ],
        column_widths_from_weights([0.65, 2.2, 2.2, 1.45]),
        font_size=9.2,
    )
    add_callout(
        doc,
        "第一条执行指令",
        "你后续可以直接对我说：“请帮我检查电脑开发环境，并初始化这个开店预测软件的项目结构。”我会从当前文件夹开始创建项目，不会一次性写完整系统。",
        fill=PALE_YELLOW,
    )


def add_phase_one(doc):
    add_heading(doc, "五、阶段1：产品原型与需求冻结（第1-2周）", 1)
    add_heading(doc, "第1周：把想法变成可开发需求", 2)
    for item in [
        "写清楚目标用户：准备开店的人、已有门店老板、加盟品牌拓展人员，MVP优先服务“准备开店的人”。",
        "确定首版输入字段：城市、详细地址或地图点、开店品类、预算区间、租金、预计客单价、营业时间、是否做外卖。",
        "确定首版输出：总评分、推荐/谨慎/不建议、客流潜力、竞争强度、成本压力、选址风险、下一步建议。",
        "确定不做内容：不做支付、不做复杂账号、不做小程序、不做自动抓美团、不承诺投资收益。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "第2周：做低保真原型", 2)
    add_table(
        doc,
        ["页面", "必须包含", "验收标准"],
        [
            ("首页/输入页", "地址输入、地图选点、品类选择、预算/租金可选输入", "用户不用看说明也知道怎么开始"),
            ("加载页", "展示正在分析周边商圈、竞品、交通和住宅", "等待过程不空白，失败时有重试提示"),
            ("报告页", "评分、分项指标、地图周边摘要、风险、建议", "报告能截图发给别人看懂"),
            ("历史页（可延后）", "保存最近的分析记录", "MVP早期可先不做登录，只保存本地或数据库记录"),
        ],
        column_widths_from_weights([1.1, 3.1, 2.3]),
    )
    add_heading(doc, "阶段1交付物", 2)
    for item in [
        "一份简短PRD：目标用户、核心场景、输入字段、输出报告结构。",
        "一张页面流程图：输入页 -> 数据分析 -> 报告页 -> 修改参数再分析。",
        "一个能点击的静态原型：可以先用HTML或React做，不需要接真实数据。",
    ]:
        add_bullet(doc, item)


def add_phase_two(doc):
    add_heading(doc, "六、阶段2：数据采集闭环（第3-4周）", 1)
    add_para(doc, "这一阶段的目标不是预测准确，而是让系统真的能根据一个地址拿到周边商业环境数据。")
    add_heading(doc, "第3周：地址与地图数据", 2)
    add_table(
        doc,
        ["任务", "实现方式", "输出字段", "验收标准"],
        [
            ("地址转经纬度", "调用高德地理编码API", "lng、lat、formatted_address", "输入一个地址能得到经纬度"),
            ("周边POI搜索", "按半径500m/1km/3km调用周边搜索", "名称、类型、距离、地址", "能列出周边商家、住宅、交通设施"),
            ("POI分类", "把高德类型映射为竞品、互补业态、住宅、办公、学校、交通、商圈", "category_group", "同品类商家能被识别为竞品"),
            ("失败处理", "API错误、Key过期、无结果、限流", "error_code、retryable", "接口失败时页面给出可理解提示"),
        ],
        column_widths_from_weights([1.2, 2.1, 1.6, 1.6]),
        font_size=9.2,
    )
    add_heading(doc, "第4周：数据存储与调试工具", 2)
    for item in [
        "建立数据库表：locations、poi_snapshots、poi_items、analysis_jobs、reports。",
        "每次分析都保存原始API返回，便于后面复盘和调参。",
        "做一个内部调试页：输入经纬度后显示原始POI、分类结果和统计摘要。",
        "先不要抓取美团/大众点评，避免早期被反爬和合规问题拖住。",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "数据闭环验收",
        "拿3个真实地址测试：一个商圈、一个社区、一个偏远位置。系统应能输出POI数量、同品类竞品数量、最近竞品距离、住宅/办公/交通设施数量。",
    )


def add_phase_three(doc):
    add_heading(doc, "七、阶段3：基础评估模型（第5-6周）", 1)
    add_para(doc, "MVP阶段建议采用可解释模型，不要一开始训练机器学习。用户需要知道“为什么推荐/为什么不推荐”，而不是只看到一个神秘分数。")
    add_heading(doc, "评分框架", 2)
    add_table(
        doc,
        ["指标", "含义", "初始权重", "计算思路"],
        [
            ("客流潜力", "周边住宅、办公、学校、交通节点和商圈活跃度", "25%", "住宅/办公/交通POI加权计分，后期可接热力数据"),
            ("竞争压力", "同品类商家数量、距离、密度", "25%", "竞品越多扣分；近距离强竞品额外扣分"),
            ("消费匹配", "品类与周边人群/业态是否匹配", "15%", "例如办公区适合快餐/咖啡，社区适合生鲜/美甲"),
            ("成本压力", "租金、人力、原料、平台费用", "15%", "用户输入或使用区域均值估算"),
            ("交通可达", "地铁、公交、停车、主干道", "10%", "交通设施数量与距离加权"),
            ("经营风险", "施工、政策、淡旺季、过度依赖单一客群", "10%", "MVP先用规则项，后续接外部数据"),
        ],
        column_widths_from_weights([1.05, 2.1, 0.75, 2.6]),
        font_size=9,
    )
    add_heading(doc, "第5周：规则评分", 2)
    for item in [
        "实现0-100总分：每个分项先归一化到0-100，再按权重求和。",
        "输出等级：80分以上推荐，60-79谨慎推荐，40-59高风险，40以下不建议。",
        "每个分项必须有解释文本，例如“500米内有12家同类奶茶店，竞争压力偏高”。",
        "所有权重写到配置文件里，不要写死在代码中，方便后面调参。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "第6周：IRS与哈夫模型雏形", 2)
    for item in [
        "IRS零售饱和指数：用需求估计/供给估计判断区域是否饱和。MVP可用住宅/办公/学校POI作为需求代理。",
        "哈夫模型：估算用户选择目标店铺而非竞品的概率，吸引力可由面积、评分、品牌、距离等因素构成；MVP阶段先用距离和竞品数量近似。",
        "不要把公式包装成绝对预测，要在报告中标明“基于公开数据和假设的模拟”。",
    ]:
        add_bullet(doc, item)


def add_phase_four(doc):
    add_heading(doc, "八、阶段4：报告生成与前端闭环（第7-8周）", 1)
    add_para(doc, "此阶段要把计算结果变成普通创业者能看懂、能行动的报告。报告是这个产品的第一价值载体。")
    add_table(
        doc,
        ["报告模块", "内容", "生成方式", "验收标准"],
        [
            ("一句话结论", "适合/谨慎/不建议开店，并说明主因", "模板+LLM润色", "10秒内看懂核心结论"),
            ("分项评分", "客流、竞争、消费匹配、成本、交通、风险", "模型直接输出", "每项有数字和解释"),
            ("周边摘要", "500m/1km内竞品、住宅、办公、交通、商圈概况", "POI统计", "能支撑结论"),
            ("经营建议", "选址、品类、定价、营业时间、外卖、营销建议", "规则库+LLM", "建议可执行，不空泛"),
            ("风险提示", "数据不完整、竞争强、成本高、模型置信度低", "规则生成", "不做投资保证"),
            ("下一步动作", "实地踩点、租金核实、竞品观察、样本访谈", "固定清单", "用户知道下一步做什么"),
        ],
        column_widths_from_weights([1.1, 1.8, 1.3, 2.3]),
        font_size=9.1,
    )
    add_heading(doc, "LLM提示词原则", 2)
    for item in [
        "把模型数值、POI统计和限制条件作为结构化JSON传给LLM，减少胡编。",
        "要求LLM只解释已有数据，不允许编造租金、客流、营业额等未提供信息。",
        "报告中固定加入免责声明：结果仅供选址参考，不构成投资建议。",
        "保留原始计算结果，LLM只负责表达，不负责最终评分。",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "MVP完成标准",
        "输入一个地址和品类后，系统能自动抓取POI、计算评分、生成报告，并在网页上展示。即使界面朴素，只要闭环真实可用，就算MVP完成。",
        fill=PALE_GREEN,
    )


def add_phase_five(doc):
    add_heading(doc, "九、阶段5：内测、回测与调参（第9-10周）", 1)
    add_para(doc, "没有真实测试，选址模型很容易看起来专业但实际无效。MVP完成后，最重要的是用真实地址反复校准。")
    add_table(
        doc,
        ["测试类型", "样本", "做法", "输出"],
        [
            ("功能测试", "10个不同地址", "检查API、页面、报告是否正常", "bug清单"),
            ("人工合理性测试", "10个你熟悉的商圈", "人工判断结果是否符合常识", "权重调整建议"),
            ("真实店铺回测", "20-50家已开店铺", "输入其地址和品类，对比现实经营状态", "模型偏差记录"),
            ("用户访谈", "5-10名潜在创业者", "让他们读报告并说出是否愿意付费", "产品修改清单"),
        ],
        column_widths_from_weights([1.2, 1.2, 2.5, 1.6]),
    )
    add_heading(doc, "调参方法", 2)
    for item in [
        "每次只调整一个权重或规则，记录调整前后的报告差异。",
        "把错误分成三类：数据错、规则错、表达错。不要用LLM润色掩盖模型错误。",
        "为每份报告保存版本号，后续能追踪是哪一版模型生成的。",
        "当用户说“不准”时，追问具体哪里不准：竞品、客流、成本、结论还是建议。",
    ]:
        add_bullet(doc, item)


def add_mirofish(doc):
    add_heading(doc, "十、MiroFish增强路线（第11-18周）", 1)
    add_para(
        doc,
        "MiroFish适合作为第二阶段差异化能力：把选址数据转换为市场环境材料，生成消费者、竞品、供应商等智能体进行模拟。建议先独立实验，再决定是否接入主系统。",
    )
    add_table(
        doc,
        ["步骤", "目标", "具体动作", "验收标准"],
        [
            ("10.1 独立部署", "确认本地能跑通MiroFish", "在单独目录克隆仓库，按README安装Node/Python依赖，使用测试材料生成一次模拟", "能得到官方示例输出"),
            ("10.2 输入改造", "把POI和评分数据转成种子材料", "生成market_context.md或JSON：区域画像、竞品、客流代理、成本假设", "MiroFish能读入你的选址材料"),
            ("10.3 智能体设计", "定义消费者、竞品、供应商、房东、平台等角色", "给每类智能体写行为规则和关注指标", "模拟结果能解释不同角色反应"),
            ("10.4 模拟约束", "控制成本和运行时间", "限制智能体数量、模拟轮次、LLM调用次数", "单次模拟在可接受成本内完成"),
            ("10.5 报告融合", "把模拟结果作为报告增强段落", "基础评分保持主结论，智能体模拟作为情景推演", "报告不因模拟随机性而前后矛盾"),
        ],
        column_widths_from_weights([1.0, 1.5, 2.8, 1.2]),
        font_size=9,
    )
    add_callout(
        doc,
        "合规提醒",
        "如果直接修改并对外提供基于AGPL v3项目的网络服务，开源义务可能影响商业化路径。上线前应核验MiroFish当前许可证和使用方式；若目标是闭源SaaS，建议保留独立重写方案和代码隔离记录。",
        fill=PALE_RED,
    )


def add_commercial(doc):
    add_heading(doc, "十一、完整商业化路线（第19-36周）", 1)
    add_table(
        doc,
        ["时间", "目标", "任务", "交付物"],
        [
            ("第19-22周", "策略模拟器", "加入调价、延长营业、做外卖、优惠活动、人力变化等变量", "经营策略模拟页面"),
            ("第23-26周", "用户体系", "登录、报告保存、套餐权限、历史对比", "用户中心与报告库"),
            ("第27-30周", "部署与性能", "云服务器、数据库迁移、缓存、日志、错误监控", "可公网访问的测试版"),
            ("第31-33周", "合规与风控", "隐私政策、用户协议、免责声明、数据留存策略、ICP备案准备", "合规文档和上线检查表"),
            ("第34-36周", "试运营", "招募种子用户、收集付费意愿、优化转化漏斗", "试运营报告和下一版路线图"),
        ],
        column_widths_from_weights([1.0, 1.2, 3.0, 1.3]),
        font_size=9.2,
    )
    add_heading(doc, "商业模式先从轻量开始", 2)
    for item in [
        "免费版：每天限定分析次数，输出基础报告。",
        "付费单份报告：提供更完整的竞品分析、策略模拟和PDF导出。",
        "订阅版：适合已有门店，持续监控周边新开/关店和经营建议。",
        "咨询服务：对加盟品牌或选址团队提供人工复核和定制数据。",
    ]:
        add_bullet(doc, item)


def add_project_management(doc):
    add_heading(doc, "十二、项目管理方式", 1)
    add_heading(doc, "建议目录结构", 2)
    add_table(
        doc,
        ["目录/文件", "用途"],
        [
            ("frontend/", "网页前端，地图选点、输入表单、报告展示"),
            ("backend/", "FastAPI后端，API路由、数据采集、模型计算、报告生成"),
            ("backend/app/services/amap.py", "高德API封装"),
            ("backend/app/services/scoring.py", "评分模型和IRS/哈夫模型"),
            ("backend/app/services/llm_report.py", "LLM提示词和报告生成"),
            ("backend/app/models/", "数据库模型"),
            ("experiments/mirofish/", "MiroFish独立实验，不与主项目代码混放"),
            ("docs/", "PRD、接口文档、测试记录、合规文档"),
        ],
        column_widths_from_weights([2.0, 4.5]),
        font_size=9.2,
    )
    add_heading(doc, "每周工作节奏", 2)
    add_table(
        doc,
        ["节奏", "动作", "结果"],
        [
            ("周初", "确定本周只做1-2个可验收目标", "减少跑偏"),
            ("开发中", "每完成一个小功能就运行测试和手动验证", "问题不过夜"),
            ("周末", "记录完成项、失败项、下周优先级", "形成项目日志"),
            ("每两周", "用真实地址跑一次完整流程", "保证不是只在假数据上好看"),
        ],
        column_widths_from_weights([1.0, 3.4, 2.1]),
    )
    add_heading(doc, "你后续可以这样让我协助", 2)
    prompts = [
        "请帮我初始化项目结构，只做前端输入页和后端健康检查接口。",
        "请帮我写高德地图POI采集脚本，并用一个地址测试。",
        "请帮我设计奶茶店选址的评分权重，并写成Python函数。",
        "请帮我把评分结果做成一份自然语言报告模板。",
        "请帮我把当前MVP跑起来，并指出下一步最该修的问题。",
        "请帮我单独研究MiroFish怎么部署，先不要合并到主项目。",
    ]
    for prompt in prompts:
        add_bullet(doc, prompt)


def add_acceptance_and_risk(doc):
    add_heading(doc, "十三、阶段验收清单", 1)
    add_table(
        doc,
        ["阶段", "必须验收", "不通过时怎么办"],
        [
            ("准备", "API Key可用、开发环境可用、首个城市和品类确定", "先不要写业务代码，补齐账号和环境"),
            ("原型", "用户能完成输入并看到报告结构", "删掉非核心页面，保留主流程"),
            ("数据", "真实地址能获取并分类POI", "检查API权限、半径、类型映射"),
            ("模型", "每个分项评分都有数据依据和解释", "先回到规则模型，不急着上ML"),
            ("报告", "报告能被非技术用户读懂并指出下一步行动", "优化结构和提示词，不让LLM编造"),
            ("内测", "至少10个真实地址跑通，记录问题", "按问题类型修复，不盲目加功能"),
            ("MiroFish", "独立实验可运行，成本和协议风险清楚", "继续隔离实验，不进入主线"),
            ("上线", "合规、部署、日志、免责声明、备份都完成", "只做小范围试运营"),
        ],
        column_widths_from_weights([1.0, 3.1, 2.4]),
        font_size=9.2,
    )
    add_heading(doc, "主要风险与控制", 1)
    add_table(
        doc,
        ["风险", "表现", "控制方式"],
        [
            ("数据成本", "高德商业授权、生活服务数据成本高", "MVP先用免费额度和用户手动补充；商业化前再评估采购"),
            ("数据合规", "抓取平台数据可能触发反爬或法律风险", "初期不抓美团；使用官方API、公开数据和用户授权数据"),
            ("模型不准", "用户认为报告与现实不符", "输出置信区间和解释，不承诺收益；持续回测"),
            ("AGPL协议", "直接改造MiroFish可能影响闭源商业化", "实验隔离、上线前法律核验、必要时独立重写"),
            ("技术复杂度", "多智能体、知识图谱、ML同时推进导致失控", "MVP只做规则模型和报告闭环"),
            ("成本失控", "LLM调用过多、模拟轮次过长", "缓存结果、限制轮次、先用小模型"),
        ],
        column_widths_from_weights([1.2, 2.3, 3.0]),
        font_size=9.2,
    )


def add_next_14_days(doc):
    add_heading(doc, "十四、从今天开始的14天行动表", 1)
    add_table(
        doc,
        ["天数", "动作", "产出", "可直接发给我的指令"],
        [
            ("第1天", "确定城市和品类，整理你最想服务的用户", "MVP范围一句话", "帮我把MVP范围写成PRD第一版"),
            ("第2天", "申请高德Key和LLM Key，安装开发工具", "账号和环境清单", "帮我检查这些Key和环境是否可用"),
            ("第3天", "初始化项目结构", "frontend/backend/docs目录", "帮我初始化项目结构并提交第一版README"),
            ("第4-5天", "做输入页原型", "地址输入、品类选择、提交按钮", "帮我做一个前端输入页原型"),
            ("第6-7天", "写高德地理编码和周边POI接口", "可运行采集脚本", "帮我写高德POI采集脚本并测试一个地址"),
            ("第8-9天", "做POI分类和统计", "竞品/住宅/交通等统计结果", "帮我把POI按开店分析类别分类"),
            ("第10-11天", "设计评分规则", "0-100分评分函数", "帮我实现奶茶店选址评分模型"),
            ("第12天", "生成报告模板", "报告JSON和自然语言文本", "帮我写LLM报告提示词"),
            ("第13天", "把前后端串起来", "网页可显示报告", "帮我联调前后端MVP流程"),
            ("第14天", "用3个真实地址测试", "问题清单和下一轮计划", "帮我分析这3份测试报告哪里需要调整"),
        ],
        column_widths_from_weights([0.8, 2.0, 1.7, 2.0]),
        font_size=8.8,
    )


def add_sources(doc):
    add_heading(doc, "十五、资料核验与注意事项", 1)
    add_para(
        doc,
        "本计划以当前文件夹中的《项目书.docx》为主要依据。涉及第三方项目许可证、地图服务费用、平台API政策等内容，在正式上线或商业化前必须重新核验。尤其是开源协议和数据抓取问题，不应只依赖项目计划书中的描述。",
    )
    add_table(
        doc,
        ["事项", "执行建议"],
        [
            ("MiroFish许可证", "在使用源码前，重新查看仓库LICENSE和README；若要商业化，保留独立重写选项并咨询法律专业人士。"),
            ("地图/POI数据", "优先使用官方API和用户授权数据，不以绕过限制的爬虫作为核心依赖。"),
            ("报告免责声明", "所有报告都应说明“仅供参考，不构成投资建议”，并展示数据来源和更新时间。"),
            ("个人信息", "只收集必要字段，地理位置和用户数据要有明确授权、删除机制和隐私政策。"),
        ],
        column_widths_from_weights([1.4, 5.1]),
        font_size=9.2,
    )


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_overview(doc)
    add_phase_zero(doc)
    add_phase_one(doc)
    add_phase_two(doc)
    add_phase_three(doc)
    add_phase_four(doc)
    add_phase_five(doc)
    add_mirofish(doc)
    add_commercial(doc)
    add_project_management(doc)
    add_acceptance_and_risk(doc)
    add_next_14_days(doc)
    add_sources(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
